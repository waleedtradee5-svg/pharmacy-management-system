import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
import json
import io
from docx import Document
from fpdf import FPDF
import os
import base64

# Ye farz kiya ja raha hai ke db_connector.py sahi se configure hai
from db_connector import fetch_data, execute_query

# Invoice attachments ke liye directory banayein agar mojood nahi hai
ATTACHMENT_DIR = "attachments"
if not os.path.exists(ATTACHMENT_DIR):
    os.makedirs(ATTACHMENT_DIR)


# --- IMAGE LOADING KO TEZ KARNE KE LIYE ---
@st.cache_data
def load_base64_image(file_path):
    """File se Base64 string load karta hai aur performance ke liye cache karta hai."""
    try:
        with open(file_path, "r") as f:
            return f.read().strip()
    except FileNotFoundError:
        st.error(f"Zaruri file nahi mil rahi: {file_path}")
        return None


# BARI STRINGS KI JAGAH AB YEH CODE ISTEMAL HOGA
BASE64_LOGO = load_base64_image(
    "C:/Users/Useless/Desktop/final erp/modules/logo_base64.txt"
)
BASE64_WATERMARK = load_base64_image(
    "C:/Users/Useless/Desktop/final erp/modules/watermark_base64.txt"
)


class SalesModule:
    """
    Sales ke poore workflow ko manage karta hai, jismein payments, returns, approvals,
    printing, aur multi-format exporting jese ERP-grade features shamil hain.
    """

    def __init__(self):
        """Module ke views aur forms ke liye session state initialize karta hai."""
        if "invoice_form_state" not in st.session_state:
            self._clear_invoice_form()

        st.session_state.setdefault("sales_view_mode", "list")
        st.session_state.setdefault("active_invoice_id", None)
        st.session_state.setdefault("current_user", "Admin User")
        st.session_state.setdefault("user_role", "Admin")
        st.session_state.setdefault("sales_filters", {"search": "", "status": "All"})

    def _clear_invoice_form(self):
        """Invoice form ko saaf state mein reset karta hai."""
        st.session_state.invoice_form_state = {
            "customer_id": None,
            "items": [{"product_id": None, "qty": 1}],
            "tax_percent": 17.0,
            "notes": "",
            "payment_method": "Card",
            "paid_amount": 0.0,
            "discount_percent": 0.0,
            "attachments": [],
            "salesperson": st.session_state.get("current_user", "Admin User"),
        }

    def _get_master_data(self):
        """Database se products aur customers ki taza tareen list haasil karta hai."""
        self.medicines = fetch_data(
            "SELECT MedicineID, MedicineName, UnitPrice, StockQty FROM medicines WHERE IsActive = TRUE"
        )
        self.customers = fetch_data(
            "SELECT id, name, phone, gender, dob, address FROM customers WHERE status = 'Active'"
        )

    def render(self):
        """Sahi view ko route karne wala main render method."""
        st.title("🧾 Sales & Invoicing")
        self._get_master_data()

        view = st.session_state.sales_view_mode
        if view == "form":
            self._render_invoice_form()
        elif view == "preview" and st.session_state.active_invoice_id:
            self._render_static_invoice_preview(st.session_state.active_invoice_id)
        else:
            self._render_invoice_list()

    def _render_invoice_list(self):
        st.subheader("Invoice Records")
        if st.button("➕ Create New Invoice", type="primary"):
            self._clear_invoice_form()
            st.session_state.sales_view_mode = "form"
            st.rerun()

        filter_cols = st.columns([2, 1])
        search = filter_cols[0].text_input(
            "Search Customer/Invoice #", st.session_state.sales_filters["search"]
        )
        status = filter_cols[1].selectbox(
            "Filter by Status",
            ["All", "Paid", "Pending", "Partially Paid", "Cancelled"],
            index=["All", "Paid", "Pending", "Partially Paid", "Cancelled"].index(
                st.session_state.sales_filters["status"]
            ),
        )

        query = "SELECT InvoiceID, InvoiceNumber, InvoiceDate, CustomerName, GrandTotal, Status FROM sales_invoices WHERE (CustomerName LIKE %s OR InvoiceNumber LIKE %s)"
        params = [f"%{search}%", f"%{search}%"]
        if status != "All":
            query += " AND Status = %s"
            params.append(status)
        query += " ORDER BY InvoiceDate DESC, InvoiceID DESC"
        invoices = fetch_data(query, tuple(params))

        if invoices.empty:
            st.info("No invoices found.")
            return

        status_colors = {
            "Paid": "green",
            "Pending": "red",
            "Partially Paid": "orange",
            "Cancelled": "gray",
        }
        for _, row in invoices.iterrows():
            with st.container(border=True):
                cols = st.columns([2, 3, 2, 2, 2])
                cols[0].markdown(
                    f"**{row['InvoiceNumber']}**<br><small>{row['InvoiceDate']}</small>",
                    unsafe_allow_html=True,
                )
                cols[1].text(row["CustomerName"])
                cols[2].text(f"Rs {row['GrandTotal']:,.2f}")
                cols[3].markdown(
                    f":{status_colors.get(row['Status'], 'gray')}[{row['Status']}]"
                )
                with cols[4]:
                    action_cols = st.columns(2)
                    if action_cols[0].button(
                        "👁️", key=f"view_{row['InvoiceID']}", help="Preview Invoice"
                    ):
                        st.session_state.active_invoice_id = row["InvoiceID"]
                        st.session_state.sales_view_mode = "preview"
                        st.rerun()
                    if action_cols[1].button(
                        "🖨️", key=f"print_{row['InvoiceID']}", help="Print Invoice"
                    ):
                        st.session_state.active_invoice_id = row["InvoiceID"]
                        st.session_state.sales_view_mode = "preview"
                        st.rerun()

    def _render_invoice_form(self):
        state = st.session_state.invoice_form_state
        st.subheader("Create New Invoice")
        customer_map = (
            self.customers.set_index("id").to_dict("index")
            if not self.customers.empty
            else {}
        )
        customer_options = {
            cid: f"{c['name']} ({c['phone']})" for cid, c in customer_map.items()
        }
        state["customer_id"] = st.selectbox(
            "Select Existing Customer",
            [None] + list(customer_options.keys()),
            format_func=lambda cid: (
                "Select..." if cid is None else customer_options[cid]
            ),
        )

        with st.expander("👤 Or Add a New Customer"):
            with st.form("new_customer_form"):
                st.markdown("**New Customer Details**")
                c_cols = st.columns(2)
                new_name = c_cols[0].text_input("Name*")
                new_phone = c_cols[1].text_input("Phone (Optional)")
                new_gender = c_cols[0].selectbox("Gender", ["Male", "Female", "Other"])
                new_dob = c_cols[1].date_input(
                    "Date of Birth", min_value=date(1920, 1, 1), max_value=date.today()
                )
                new_address = st.text_area("Address")
                if st.form_submit_button("Save Customer", type="primary"):
                    if new_name:
                        execute_query(
                            "INSERT INTO customers (name,phone,gender,dob,address,status) VALUES (%s,%s,%s,%s,%s,'Active')",
                            (
                                new_name,
                                new_phone or None,
                                new_gender,
                                new_dob,
                                new_address,
                            ),
                        )
                        st.toast("✅ Customer Added!")
                        st.rerun()
                    else:
                        st.error("Customer Name is required.")

        st.markdown("---")
        st.markdown("#### Invoice Items")
        medicine_map = (
            self.medicines.set_index("MedicineID").to_dict("index")
            if not self.medicines.empty
            else {}
        )
        subtotal = 0
        for i, item in enumerate(state["items"]):
            cols = st.columns([4, 2, 2, 2, 1])
            product_options = {
                mid: f"{m['MedicineName']} ({m['StockQty']} left)"
                for mid, m in medicine_map.items()
                if m["StockQty"] > 0 or mid == item.get("product_id")
            }
            selected_id = cols[0].selectbox(
                "Product*",
                [None] + list(product_options.keys()),
                format_func=lambda mid: (
                    "Select..." if mid is None else product_options[mid]
                ),
                key=f"product_{i}",
            )
            if selected_id and selected_id != item.get("product_id"):
                item.update(
                    {
                        "product_id": selected_id,
                        "price": float(medicine_map[selected_id]["UnitPrice"]),
                    }
                )
            item["price"] = cols[1].number_input(
                "Price", value=item.get("price", 0.0), disabled=True, key=f"price_{i}"
            )
            max_stock = medicine_map.get(item.get("product_id"), {}).get("StockQty", 0)
            item["qty"] = cols[2].number_input(
                "Quantity*",
                min_value=1,
                max_value=max_stock if max_stock > 0 else 1,
                value=item.get("qty", 1),
                key=f"qty_{i}",
            )
            line_total = item.get("qty", 1) * item.get("price", 0.0)
            subtotal += line_total
            cols[3].text_input(
                "Total", f"Rs {line_total:,.2f}", disabled=True, key=f"total_{i}"
            )
            if cols[4].button("🗑️", key=f"del_item_{i}", help="Remove item"):
                state["items"].pop(i)
                st.rerun()

        if st.button("➕ Add Another Product"):
            state["items"].append({"product_id": None, "qty": 1})
            st.rerun()

        st.markdown("---")
        final_cols = st.columns(2)
        with final_cols[0]:
            state["notes"] = st.text_area("Notes", value=state["notes"])
            state["payment_method"] = st.selectbox(
                "Payment Method",
                ["Card", "Cash", "Bank Transfer"],
                index=["Card", "Cash", "Bank Transfer"].index(state["payment_method"]),
            )
        with final_cols[1]:
            d_cols = st.columns(2)
            state["discount_percent"] = d_cols[0].number_input(
                "Discount (%)",
                min_value=0.0,
                max_value=100.0,
                value=state.get("discount_percent", 0.0),
                format="%.2f",
            )
            state["tax_percent"] = d_cols[1].number_input(
                "Tax (%)",
                min_value=0.0,
                value=state.get("tax_percent", 17.0),
                format="%.2f",
            )
            discount_amount = subtotal * (state["discount_percent"] / 100)
            subtotal_after_discount = subtotal - discount_amount
            tax_amount = subtotal_after_discount * (state["tax_percent"] / 100)
            grand_total = subtotal_after_discount + tax_amount
            st.markdown(f"### Grand Total: Rs {grand_total:,.2f}")
            state["paid_amount"] = st.number_input(
                "Amount Paid",
                min_value=0.0,
                max_value=grand_total if grand_total > 0 else 1.0,
                value=grand_total,
            )
            balance_due = grand_total - state["paid_amount"]
            st.metric(
                "Balance Due",
                f"Rs {balance_due:,.2f}",
                delta=f"-Rs {state['paid_amount']:,.2f}",
                delta_color="inverse",
            )

        st.markdown("---")
        action_cols = st.columns([2, 1])
        if action_cols[0].button(
            "✅ Save & Generate Invoice", type="primary", use_container_width=True
        ):
            invoice_id = self._save_invoice(
                subtotal, tax_amount, grand_total, balance_due, discount_amount
            )
            if invoice_id:
                st.session_state.active_invoice_id = invoice_id
                st.session_state.sales_view_mode = "preview"
                st.rerun()
        if action_cols[1].button("Clear Form", use_container_width=True):
            self._clear_invoice_form()
            st.rerun()

    def _save_invoice(
        self, subtotal, tax_amount, grand_total, balance_due, discount_amount
    ):
        state = st.session_state.invoice_form_state
        if not state["customer_id"]:
            st.error("Customer select karna lazmi hai.")
            return None
        if not all(item.get("product_id") for item in state["items"]):
            st.error("Har item line mein product select karna lazmi hai.")
            return None

        now = datetime.now()
        invoice_number = f"MUJ-{now.strftime('%Y%m%d-%H%M%S')}"
        customer_info = self.customers[
            self.customers["id"] == state["customer_id"]
        ].iloc[0]
        age = (
            (now.date() - pd.to_datetime(customer_info["dob"]).date()).days // 365
            if pd.notna(customer_info.get("dob"))
            else None
        )
        items_json = json.dumps(
            [
                {
                    "MedicineID": i["product_id"],
                    "MedicineName": self.medicines[
                        self.medicines["MedicineID"] == i["product_id"]
                    ].iloc[0]["MedicineName"],
                    "Quantity": i["qty"],
                    "UnitPrice": i["price"],
                    "LineTotal": i["qty"] * i["price"],
                }
                for i in state["items"]
            ]
        )
        status = (
            "Paid"
            if balance_due <= 0
            else "Partially Paid" if state["paid_amount"] > 0 else "Pending"
        )

        query = "INSERT INTO sales_invoices (InvoiceNumber,CustomerName,CustomerPhone,CustomerGender,CustomerAge,Salesperson,InvoiceDate,InvoiceTime,SubTotal,TaxPercent,TaxAmount,GrandTotal,ItemsData,Status,Notes,DiscountType,DiscountValue,PaidAmount,BalanceDue,ApprovalStatus) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'Pending')"
        params = (
            invoice_number,
            customer_info["name"],
            customer_info["phone"],
            customer_info.get("gender"),
            age,
            state["salesperson"],
            now.date(),
            now.time(),
            subtotal,
            state["tax_percent"],
            tax_amount,
            grand_total,
            items_json,
            status,
            state["notes"],
            "Percentage",
            state["discount_percent"],
            state["paid_amount"],
            balance_due,
        )

        success, last_id = execute_query(query, params, return_last_id=True)
        if success and last_id:
            if state["paid_amount"] > 0:
                execute_query(
                    "INSERT INTO invoice_payments (InvoiceID,Amount,PaymentMethod,PaymentDate) VALUES (%s,%s,%s,%s)",
                    (
                        last_id,
                        state["paid_amount"],
                        state["payment_method"],
                        now.date(),
                    ),
                )
            for item in state["items"]:
                execute_query(
                    "UPDATE medicines SET StockQty=StockQty-%s WHERE MedicineID=%s",
                    (item["qty"], item["product_id"]),
                )
            st.success(f"Invoice {invoice_number} save ho gaya!")
            return last_id
        else:
            st.error("Invoice save nahi ho saka. Stock aur payment update nahi hue.")
            return None

    def _render_static_invoice_preview(self, invoice_id):
        try:
            invoice_data = fetch_data(
                "SELECT * FROM sales_invoices WHERE InvoiceID=%s", (invoice_id,)
            ).iloc[0]
            items_df = pd.DataFrame(json.loads(invoice_data["ItemsData"]))
        except (IndexError, KeyError):
            st.error("Invoice data nahi mil saka.")
            st.session_state.sales_view_mode = "list"
            st.rerun()
            return

        # NEW: Yeh naya aur behtar print ka tareeqa hai
        self._render_print_component(invoice_data, items_df)

        st.markdown("<hr class='no-print'>", unsafe_allow_html=True)
        st.markdown(
            self._generate_invoice_html(invoice_data, items_df), unsafe_allow_html=True
        )
        st.markdown(self._get_invoice_styles(), unsafe_allow_html=True)

    def _render_print_component(self, invoice_data, items_df):
        """
        Ek self-contained HTML component banata hai jismein print button aur script dono shamil hain.
        Yeh tareeqa sab se reliable hai.
        """
        full_html_page = self._generate_printable_html_page(invoice_data, items_df)
        html_for_js = json.dumps(full_html_page)

        component_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ margin: 0; padding: 0; }}
                button {{
                    background-color: #0068c9; color: white; border: none;
                    padding: 10px 24px; text-align: center; text-decoration: none;
                    display: inline-block; font-size: 16px; margin: 4px 2px;
                    cursor: pointer; border-radius: 8px; width: 100%;
                }}
                button:hover {{ opacity: 0.9; }}
            </style>
        </head>
        <body>
            <div style="display: flex; justify-content: space-between; gap: 10px;">
                <button id="print-button">🖨️ Print Invoice</button>
            </div>
            <script>
                function printInvoice(htmlContent) {{
                    let iframe = window.parent.document.getElementById('printing-iframe');
                    if (!iframe) {{
                        iframe = window.parent.document.createElement('iframe');
                        iframe.id = 'printing-iframe';
                        iframe.style.display = 'none';
                        window.parent.document.body.appendChild(iframe);
                    }}
                    const doc = iframe.contentWindow.document;
                    doc.open();
                    doc.write(htmlContent);
                    doc.close();
                    
                    setTimeout(function() {{
                        iframe.contentWindow.focus();
                        iframe.contentWindow.print();
                    }}, 500);
                }}

                document.getElementById('print-button').addEventListener('click', function() {{
                    printInvoice({html_for_js});
                }});
            </script>
        </body>
        </html>
        """

        st.markdown("<div class='no-print action-bar'>", unsafe_allow_html=True)
        st.subheader("Invoice Actions")

        cols = st.columns([1.5, 2, 2])
        if cols[0].button("⬅️ Back to List", use_container_width=True):
            st.session_state.sales_view_mode = "list"
            st.session_state.active_invoice_id = None
            st.rerun()

        with cols[1]:
            st.components.v1.html(component_html, height=50)

        with cols[2].popover("📥 Export As", use_container_width=True):
            st.download_button(
                "📕 PDF",
                self._generate_pdf_manually(invoice_data, items_df),
                f"{invoice_data['InvoiceNumber']}.pdf",
                "application/pdf",
                use_container_width=True,
            )
            # Add other export buttons here

        st.markdown("</div>", unsafe_allow_html=True)

    def _generate_printable_html_page(self, invoice_data, items_df):
        """Printing ke liye ek mukammal, standalone HTML page banata hai."""
        invoice_body = self._generate_invoice_html(invoice_data, items_df)
        styles = self._get_invoice_styles()

        return f"""
        <!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Invoice {invoice_data['InvoiceNumber']}</title>{styles}</head><body>{invoice_body}</body></html>
        """

    def _generate_invoice_html(self, invoice_data, items_df):
        logo_data_uri = f"data:image/png;base64,{BASE64_LOGO}" if BASE64_LOGO else ""
        watermark_data_uri = (
            f"data:image/png;base64,{BASE64_WATERMARK}" if BASE64_WATERMARK else ""
        )
        time_str = "N/A"
        if pd.notna(invoice_data["InvoiceTime"]):
            time_val = invoice_data["InvoiceTime"]
            time_obj = (
                (datetime.min + time_val).time()
                if isinstance(time_val, timedelta)
                else None
            )
            if time_obj:
                time_str = time_obj.strftime("%I:%M %p")
        items_html = "".join(
            [
                f"<tr class='item'><td class='item-name'>{row['MedicineName']}</td><td class='qty'>{row['Quantity']}</td><td class='price'>Rs {row['UnitPrice']:,.2f}</td><td class='total'>Rs {row['LineTotal']:,.2f}</td></tr>"
                for _, row in items_df.iterrows()
            ]
        )
        balance_due_class = (
            "balance-due-red"
            if float(invoice_data.get("BalanceDue", 0)) > 0
            else "balance-due-green"
        )
        subtotal = float(invoice_data["SubTotal"])
        discount_percent = float(invoice_data.get("DiscountValue", 0))
        discount_amount = subtotal * (discount_percent / 100)
        return f"""<div class="invoice-container" style="--watermark-url: url('{watermark_data_uri}');"><div class="header-section"><div class="company-logo"><img src="{logo_data_uri}" alt="Logo"></div><div class="company-info"><h2>Mujtabah Pharmacy</h2><p>Model Town, Lahore</p><p>+92 333 1234567 | info@mujtabapharmacy.com</p><p>GST#: 12-345678-9</p></div></div><div class="invoice-title">INVOICE</div><div class="details-section"><div class="bill-to-info"><strong>Bill To:</strong><br>{invoice_data['CustomerName']}<br>{invoice_data.get('CustomerAddress', 'N/A')}<br>{invoice_data['CustomerPhone'] or 'N/A'}</div><div class="invoice-meta-info"><strong>Invoice #:</strong> {invoice_data['InvoiceNumber']}<br><strong>Date:</strong> {pd.to_datetime(invoice_data['InvoiceDate']).strftime('%B %d, %Y')}<br><strong>Time:</strong> {time_str}<br><strong>Due Date:</strong> {(pd.to_datetime(invoice_data['InvoiceDate']) + timedelta(days=15)).strftime('%B %d, %Y')}</div></div><table class="items-table"><thead><tr class="heading"><th class="item-name">Item & Description</th><th class="qty">Qty</th><th class="price">Unit Price</th><th class="total">Line Total</th></tr></thead><tbody>{items_html}</tbody></table><div class="totals-summary"><table class="totals-table"><tr><td>Subtotal:</td><td class="right">Rs {subtotal:,.2f}</td></tr><tr><td>Discount ({discount_percent}%):</td><td class="right">-Rs {discount_amount:,.2f}</td></tr><tr><td>Tax ({invoice_data['TaxPercent']}%):</td><td class="right">Rs {invoice_data['TaxAmount']:,.2f}</td></tr><tr class="grand-total"><td>Grand Total:</td><td class="right">Rs {invoice_data['GrandTotal']:,.2f}</td></tr><tr><td>Amount Paid:</td><td class="right">Rs {invoice_data.get('PaidAmount', 0):,.2f}</td></tr><tr class="{balance_due_class}"><td>Balance Due:</td><td class="right">Rs {invoice_data.get('BalanceDue', 0):,.2f}</td></tr></table></div><div class="footer-section"><p><strong>Payment Method:</strong> {invoice_data.get('PaymentMethod', 'N/A')}</p><p><strong>Notes:</strong> {invoice_data['Notes'] or 'Thank you for your business!'}</p><p><strong>Terms & Conditions:</strong> Payment is due within 15 days of the invoice date.</p><hr><p class="website-info">www.mujtapharmacy.com</p></div></div>"""

    def _get_invoice_styles(self):
        return """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
    :root {
        --primary-blue: #2196F3;
        --light-gray: #f5f5f5;
        --medium-gray: #e0e0e0;
        --dark-gray-text: #333;
        --light-text: #666;
        --red-alert: #dc3545;
        --green-success: #28a745;
    }
    .invoice-container {
        position: relative;
        max-width: 850px;
        margin: auto;
        padding: 40px;
        border: 1px solid var(--medium-gray);
        box-shadow: 0 5px 20px rgba(0,0,0,.1);
        font-family: 'Roboto', sans-serif;
        color: var(--dark-gray-text);
        background: #fff;
        line-height: 1.6;
        animation: fadeIn .5s ease-in-out;
    }
    .invoice-container::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background-image: var(--watermark-url);
        background-position: center;
        background-repeat: no-repeat;
        background-size: contain;
        opacity: 0.08;
        z-index: 0;
    }
    .header-section {
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        margin-bottom: 30px;
        padding-bottom: 20px;
        border-bottom: 2px solid var(--medium-gray);
        position: relative;
        z-index: 1;
    }
    .company-logo img {
        width: 120px;
        height: auto;
    }
    .company-info {
        text-align: right;
    }
    .company-info h2 {
        color: var(--primary-blue);
        margin: 0 0 5px 0;
        font-size: 1.8em;
        font-weight: 700;
    }
    .company-info p {
        margin: 0;
        font-size: .9em;
        color: var(--light-text);
    }
    .invoice-title {
        text-align: center;
        font-size: 3em;
        color: var(--primary-blue);
        margin-bottom: 30px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 2px;
        position: relative;
        z-index: 1;
    }
    .details-section {
        display: flex;
        justify-content: space-between;
        margin-bottom: 40px;
        padding: 15px 0;
        border-bottom: 1px solid var(--light-gray);
        position: relative;
        z-index: 1;
    }
    .bill-to-info, .invoice-meta-info {
        width: 48%;
        font-size: .95em;
    }
    .invoice-meta-info {
        text-align: right;
    }
    .items-table {
        width: 100%;
        border-collapse: collapse;
        margin-bottom: 30px;
        position: relative;
        z-index: 1;
    }
    .items-table thead tr.heading {
        background: var(--primary-blue);
        color: #fff;
    }
    .items-table th, .items-table td {
        padding: 12px 15px;
        text-align: left;
        border-bottom: 1px solid var(--light-gray);
    }
    .totals-summary {
        display: flex;
        justify-content: flex-end;
        margin-top: 20px;
        padding-top: 20px;
        border-top: 2px solid var(--medium-gray);
        position: relative;
        z-index: 1;
    }
    .totals-table {
        width: 45%;
    }
    .totals-table td {
        padding: 8px 15px;
    }
    .totals-table tr.grand-total td {
        font-weight: 700;
        font-size: 1.3em;
        color: var(--primary-blue);
        border-top: 1px solid var(--medium-gray);
    }
    .totals-table tr.balance-due-red td {
        font-weight: 700;
        font-size: 1.2em;
        color: var(--red-alert);
    }
    .totals-table tr.balance-due-green td {
        font-weight: 700;
        font-size: 1.2em;
        color: var(--green-success);
    }
    @media print {
        /* Yeh naya code hai jo page size aur margins set karega */
        @page {
            size: A4 portrait;
            margin: 10mm;
        }

        .no-print {
            display: none !important;
        }
        .main .block-container {
            padding: 0 !important;
            margin: 0 !important;
        }
        .invoice-container {
            box-shadow: none;
            border: none;
            margin: 0;
            max-width: 100%;
        }
        body {
            -webkit-print-color-adjust: exact !important;
            print-color-adjust: exact !important;
        }
        .invoice-container::before {
            opacity: 0.08 !important;
        }
    }
    </style>
    """
    # --- EXPORT AND PDF FUNCTIONS ---
    def _export_to_csv(self, items_df):
        return items_df.to_csv(index=False).encode("utf-8")

    def _export_to_json(self, invoice_data, items_df):
        return json.dumps(
            {
                "invoice_details": invoice_data.to_dict(),
                "items": items_df.to_dict("records"),
            },
            indent=4,
            default=str,
        ).encode("utf-8")

    def _export_to_txt(self, invoice_data, items_df):
        txt = f"INVOICE - Mujtabah Pharmacy\n----------------------------------\nInvoice #: {invoice_data['InvoiceNumber']}\nDate: {pd.to_datetime(invoice_data['InvoiceDate']).strftime('%Y-%m-%d')}\nCustomer: {invoice_data['CustomerName']}\n----------------------------------\nItems:\n"
        for _, row in items_df.iterrows():
            txt += f"- {row['MedicineName']} (Qty: {row['Quantity']}) @ Rs {row['UnitPrice']:,.2f} = Rs {row['LineTotal']:,.2f}\n"
        txt += f"----------------------------------\nSubtotal: Rs {invoice_data['SubTotal']:,.2f}\nGrand Total: Rs {invoice_data['GrandTotal']:,.2f}"
        return txt.encode("utf-8")

    def _export_to_excel(self, invoice_data, items_df):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            pd.DataFrame([invoice_data.drop("ItemsData")]).T.set_axis(
                ["Details"], axis=1
            ).to_excel(writer, sheet_name="Invoice Summary")
            items_df.to_excel(writer, sheet_name="Items", index=False)
        return output.getvalue()

    def _export_to_word(self, invoice_data, items_df):
        doc = Document()
        doc.add_heading(f"Invoice: {invoice_data['InvoiceNumber']}", 0)
        doc.add_paragraph(
            f"Date: {pd.to_datetime(invoice_data['InvoiceDate']).strftime('%B %d, %Y')}\nCustomer: {invoice_data['CustomerName']}"
        )
        doc.add_heading("Items", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = (
            "Item Name",
            "Quantity",
            "Unit Price",
            "Total",
        )
        for _, row in items_df.iterrows():
            row_cells = table.add_row().cells
            (
                row_cells[0].text,
                row_cells[1].text,
                row_cells[2].text,
                row_cells[3].text,
            ) = (
                row["MedicineName"],
                str(row["Quantity"]),
                f"{row['UnitPrice']:.2f}",
                f"{row['LineTotal']:.2f}",
            )
        doc.add_paragraph(
            f"\nSubtotal: Rs {invoice_data['SubTotal']:,.2f}\nGrand Total: Rs {invoice_data['GrandTotal']:,.2f}"
        )
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def _generate_pdf_manually(self, invoice_data, items_df):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        # (Yahan aapka poora PDF generation ka code paste karein, maine isay chota rakha hai)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, f"Invoice: {invoice_data['InvoiceNumber']}", 0, 1, "C")
        return bytes(pdf.output())
